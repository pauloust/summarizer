#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Wed Jan 29 15:04:14 2020

@author: paul
"""


import numpy as np
import pandas as pd
import nltk
import networkx as nx

from nltk.tokenize import sent_tokenize
from nltk.corpus import stopwords

from sklearn.metrics.pairwise import cosine_similarity

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_ALIGN_PARAGRAPH

nltk.download('stopwords')
stop_words = stopwords.words('english')


def xls_to_csv(path):
    df = pd.read_excel(path) 
    df.to_csv("./Cluster.csv", sep=";")
    return(df)

def articles(path):
    articles = []
    df = xls_to_csv(path)
    for i in range(len(df)):
        articles.append(df['text'][i])
    return(articles)
    
def sentences(path):
    df = xls_to_csv(path)
    sent = []
    for s in df['text']:
        sent.append(sent_tokenize(s))
    sent = [y for x in sent for y in x]
    return(sent)
    
def remove_stopwords(sentences):
    sen_new = " ".join([i for i in sentences if i not in stop_words])
    return(sen_new)

def clean_sentences(path):
    sent = sentences(path)
    clean_sentences = pd.Series(sent).str.replace("[^a-zA-Z]", " ")
    clean_sentences = [s.lower() for s in clean_sentences]
    clean_sentences = [remove_stopwords(r.split()) for r in clean_sentences]
    return(clean_sentences)
    
def glove_matrix():
    matrice_vector = 'glove.6B.300d.txt'
    word_embeddings = {}
    f = open(matrice_vector, encoding='utf-8')
    for line in f:
        values = line.split()
        word = values[0]
        coefs = np.asarray(values[1:], dtype='float32')
        word_embeddings[word] = coefs
    f.close()
    return(word_embeddings)

word_embeddings = glove_matrix()
n = len(word_embeddings.get('car'))

def sentence_vectors(path):
    sent_vect = []
    clean_sent = clean_sentences(path)
    for i in clean_sent: 
        if len(i) != 0:
            v = sum([word_embeddings.get(w, np.zeros((n,))) for w in i.split()])/(len(i.split()))
        else:
            v = np.zeros((n,))
        sent_vect.append(v)
    return(sent_vect)

def similarity_matrix(path):
    m = len(sentences(path))
    sent_vect = sentence_vectors(path)
    sim_mat = np.zeros([m, m])
    for i in range(m):
        for j in range(m):
            if i != j:
                sim_mat[i][j] = cosine_similarity(sent_vect[i].reshape(1,n),sent_vect[j].reshape(1,n))[0,0]  
    return(sim_mat)

def scores(path):
    sim_mat = similarity_matrix(path)
    nx_graph = nx.from_numpy_array(sim_mat)
    s = nx.pagerank(nx_graph)
    return(s)

def rank(path):
    scs = scores(path)
    sent = sentences(path)
    ranked_sentences = sorted(((scs[i],s) for i,s in enumerate(sent)), reverse=True)
    return(ranked_sentences)

def glove_simple(path, nb_phrases = 10):
    ranked_sentences = rank(path)
    temp = ranked_sentences[0][1]
    for i in range(1,nb_phrases):
        temp = temp + ' ' + ranked_sentences[i][1]
    return(temp)
    
def export_to_word(path, name, nb_phrases = 10):
    articles_list = articles(path)
    articles_lists = []
    for art in articles_list:
        articles_lists.append(sent_tokenize(art))
    
    resume = glove_simple(path, nb_phrases)
    
    rank_list = []
    for s in sent_tokenize(resume):
        rank_list.append(s)
    
    titles_list = []
    df = xls_to_csv(path)
    for i in range(len(df)):
        titles_list.append(df['Title'][i])
    
    document = Document()
    
    document.add_heading(name, 0)
    
    document.add_heading('GloVe Simple Summary', level=1)
    resume = document.add_paragraph(resume)
    resume.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    error1 = resume.text.find(').')
    while error1 > -1 and resume.text[error1+2] != ' ':
        pre = resume.text[:error1+2]
        post = resume.text[error1+2:]
        error1 = post.find(').')
        resume.text = pre
        resume.add_run(' ')
        resume.add_run(post)
    error2 = resume.text.find(' .')
    while error2 > -1:
        pre = resume.text[:error2]
        post = resume.text[error2+2:]
        resume.text = pre
        resume.add_run('. ')
        resume.add_run(post)
        error2 = resume.text.find(' .')
    
    for i in range(len(articles_lists)):
        document.add_heading(titles_list[i], level=1)
        para = document.add_paragraph()
        for j in range(len(articles_lists[i])):
            resume = sent_tokenize(document.paragraphs[2].text)
            if articles_lists[i][j] in resume:
                run = para.add_run(articles_lists[i][j])
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                para.add_run(' ')
            else : 
                para.add_run(articles_lists[i][j])
                para.add_run(' ')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            
    file_name = name + '.docx'
    document.save(file_name)




