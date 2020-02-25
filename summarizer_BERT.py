#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Thu Dec 19 11:38:39 2019

@author: paul
"""

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_ALIGN_PARAGRAPH

from summarizer import Summarizer
import pandas as pd
from nltk.tokenize import sent_tokenize

def xls_to_csv(path):
    df = pd.read_excel(path) 
    df.to_csv("./Cluster.csv", sep=";")
    return(df)

def articles(path):
    articles = []
    df = xls_to_csv(path)
    for i in range(len(df)):
        articles.append(df['text'][i])
    return(articles)

def sentences(path):
    df = xls_to_csv(path)
    sent = []
    for s in df['text']:
        sent.append(sent_tokenize(s))
    sent = [y for x in sent for y in x]
    return(sent)

def BERT(path, ratio):
    text = ''
    sent = sentences(path)
    
    for i in range(len(sent)):
        text = text + sent[i]
    
    model = Summarizer()
    result = model(text, ratio = ratio, min_length=60)

    resume = sent_tokenize(result)[0]
    for s in sent_tokenize(result)[1:]:
        resume = resume + ' ' + s
    return(resume)

def cluster_summarizer_BERT(path, name, ratio=0.15):
    articles_list = articles(path)
    articles_lists = []
    for art in articles_list:
        articles_lists.append(sent_tokenize(art))
    
    summary = BERT(path, ratio)
    
    rank_list = []
    for s in sent_tokenize(summary):
        rank_list.append(s)
    
    titles_list = []
    df = xls_to_csv(path)
    for i in range(len(df)):
        titles_list.append(df['Title'][i])
    
    document = Document()
    
    document.add_heading(name, 0)
    
    document.add_heading('BERT Summary', level=1)
    resume = document.add_paragraph(summary)
    resume.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    error1 = resume.text.find(').')
    while error1 > -1 and error1+2 <= len(resume.text)-1 and resume.text[error1+2] != ' ':
        pre = resume.text[:error1+2]
        post = resume.text[error1+2:]
        error1 = post.find(').')
        resume.text = pre
        resume.add_run(' ')
        resume.add_run(post)
    error2 = resume.text.find(' .')
    while error2 > -1:
        pre = resume.text[:error2]
        post = resume.text[error2+2:]
        resume.text = pre
        resume.add_run('. ')
        resume.add_run(post)
        error2 = resume.text.find(' .')
    
    for i in range(len(articles_lists)):
        document.add_heading(titles_list[i], level=1)
        para = document.add_paragraph()
        for j in range(len(articles_lists[i])):
            resume = sent_tokenize(document.paragraphs[2].text)
            if articles_lists[i][j] in resume:
                run = para.add_run(articles_lists[i][j])
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                para.add_run(' ')
            else : 
                para.add_run(articles_lists[i][j])
                para.add_run(' ')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            
    file_name = name + '.docx'
    document.save(file_name)


def article_summarizer_BERT(path, name, ratio):
    
    articles_list = articles(path)
    articles_lists = []
    for art in articles_list:
        articles_lists.append(sent_tokenize(art))
    
    model = Summarizer()
    
    summaries = []
    
    for x in articles_list :
        result = model(x, ratio = ratio, min_length=60)
        summaries.append(result)
    
    titles_list = []
    df = xls_to_csv(path)
    for i in range(len(df)):
        titles_list.append(df['Title'][i])
    
    document = Document()
    
    document.add_heading(name, 0)
    
    for i in range(len(summaries)):
        
        document.add_heading('Résumé : '+titles_list[i], level=1)
        resume = document.add_paragraph(summaries[i])
        resume.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        error1 = resume.text.find(').')
        while error1 > -1 and error1+2 <= len(resume.text)-1 and resume.text[error1+2] != ' ':
            pre = resume.text[:error1+2]
            post = resume.text[error1+2:]
            error1 = post.find(').')
            resume.text = pre
            resume.add_run(' ')
            resume.add_run(post)
        error2 = resume.text.find(' .')
        while error2 > -1:
            pre = resume.text[:error2]
            post = resume.text[error2+2:]
            resume.text = pre
            resume.add_run('. ')
            resume.add_run(post)
            error2 = resume.text.find(' .')
        
        document.add_heading(titles_list[i], level=1)
        para = document.add_paragraph()
        for j in range(len(articles_lists[i])):
            resume = sent_tokenize(document.paragraphs[2].text)
            if articles_lists[i][j] in summaries[i]:
                run = para.add_run(articles_lists[i][j])
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                para.add_run(' ')
            else : 
                para.add_run(articles_lists[i][j])
                para.add_run(' ')
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
    file_name = name + '.docx'
    document.save(file_name)
    
    
    
    